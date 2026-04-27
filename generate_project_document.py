from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_mono_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def main():
    doc = Document()

    title = doc.add_heading("Serverless File Upload System - Project Documentation", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph("Prepared for project explanation, architecture review, and presentation.")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    add_heading(doc, "1. Project Overview", level=1)
    doc.add_paragraph(
        "This project is a cloud-native Serverless File Upload System. Users upload files from a React frontend. "
        "The backend is powered by AWS services: API Gateway, Lambda, S3, and DynamoDB."
    )
    doc.add_paragraph(
        "The main goal is to provide scalable, low-maintenance file management where compute resources are used only when needed."
    )

    add_heading(doc, "2. Technology Stack", level=1)
    stack = [
        "Frontend: React + Vite + Tailwind CSS",
        "API Layer: AWS API Gateway (HTTP API)",
        "Compute: AWS Lambda (Node.js)",
        "Storage: Amazon S3 (file objects)",
        "Metadata DB: Amazon DynamoDB",
        "Local Development: Vite proxy for API integration",
    ]
    for item in stack:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "3. Architecture Diagram", level=1)
    doc.add_paragraph("Logical architecture of the system:")
    architecture = r"""
+-----------------------+        HTTPS         +-----------------------+
|  React Frontend UI    | ------------------>  |   API Gateway (HTTP)  |
|  (Login/Dashboard/    |                      |  /dashboard /files    |
|   Upload/My Files)    | <------------------  |  /files/{id}          |
+-----------------------+      JSON Response   +-----------+-----------+
                                                         |
                                                         v
                                              +-----------------------+
                                              |   AWS Lambda Handler  |
                                              |  (Route-based logic)  |
                                              +-----+-----------+-----+
                                                    |           |
                                 Put/Get/Delete     |           | Scan/Get/Put/Delete
                                                    v           v
                                           +-------------+   +----------------+
                                           | Amazon S3   |   |  DynamoDB      |
                                           | File Object |   | File Metadata  |
                                           +-------------+   +----------------+
"""
    add_mono_block(doc, architecture.strip("\n"))

    add_heading(doc, "4. Flow Diagram (End-to-End Request Flow)", level=1)
    doc.add_paragraph("File upload and retrieval flow:")
    flow = r"""
[User Selects File]
        |
        v
[Upload Page Validation]
  - type check
  - size check
        |
        v
[Upload Success UI + S3 Key generated]
        |
        v
[POST /files to API Gateway]
        |
        v
[Lambda: parse body + create metadata item]
        |
        v
[DynamoDB: PutItem]
        |
        v
[GET /files from My Files page]
        |
        v
[Lambda: Scan/Query DynamoDB]
        |
        v
[Frontend renders file cards/table]
"""
    add_mono_block(doc, flow.strip("\n"))

    add_heading(doc, "5. How the Project Works", level=1)
    add_heading(doc, "5.1 Frontend Layer", level=2)
    frontend_points = [
        "User authentication screen provides login/register flow with animated UI.",
        "Dashboard page loads aggregated stats and recent uploads from backend APIs.",
        "Upload page validates selected files (type and size), shows progress animation, and submits metadata.",
        "My Files page supports search, filter, sort, pagination, and file detail navigation.",
        "Settings page manages profile info, email updates, and password reset UI.",
    ]
    for p in frontend_points:
        doc.add_paragraph(p, style="List Bullet")

    add_heading(doc, "5.2 API Layer (API Gateway + Lambda)", level=2)
    api_points = [
        "API Gateway exposes routes such as /dashboard, /files, /files/{id}.",
        "Lambda receives method + path and routes requests to the correct handler logic.",
        "CORS headers are included so browser-based frontend can call API safely.",
    ]
    for p in api_points:
        doc.add_paragraph(p, style="List Bullet")

    add_heading(doc, "5.3 Data Layer", level=2)
    data_points = [
        "S3 stores uploaded file objects (binary files).",
        "DynamoDB stores metadata: fileId, fileName, fileType, fileSizeBytes, uploadedAt, uploadedBy, status, s3Key.",
        "Frontend views are driven primarily by metadata returned from DynamoDB through Lambda.",
    ]
    for p in data_points:
        doc.add_paragraph(p, style="List Bullet")

    add_heading(doc, "6. Core Backend Routes", level=1)
    routes = [
        "GET /dashboard -> Returns { totalFiles, storageUsed, lastUpload }",
        "GET /files -> Returns file metadata list",
        "POST /files -> Inserts a new file metadata record",
        "GET /files/{id} -> Returns metadata of one file",
        "DELETE /files/{id} -> Deletes one metadata record",
    ]
    for r in routes:
        doc.add_paragraph(r, style="List Bullet")

    add_heading(doc, "7. Security and Best Practices", level=1)
    security = [
        "IAM user (not root) is used for daily access.",
        "MFA enabled for secure account authentication.",
        "Least-privilege IAM policies recommended for Lambda and developers.",
        "S3 public access should remain blocked unless explicitly required.",
        "Rotate exposed access keys immediately if accidentally shared.",
    ]
    for s in security:
        doc.add_paragraph(s, style="List Bullet")

    add_heading(doc, "8. Deployment and Testing Checklist", level=1)
    checklist = [
        "Deploy Lambda code and confirm correct handler file is configured.",
        "Attach Lambda integration to all API Gateway routes.",
        "Set route authorization to NONE for public testing routes.",
        "Deploy API Gateway stage (prod) after any route change.",
        "Verify endpoints return JSON (not default Hello from Lambda).",
        "Test frontend integration via /api proxy in local development.",
    ]
    for c in checklist:
        doc.add_paragraph(c, style="List Bullet")

    add_heading(doc, "9. Conclusion", level=1)
    doc.add_paragraph(
        "This project demonstrates a practical implementation of modern serverless architecture. "
        "It combines scalable storage (S3), event-driven compute (Lambda), API management (API Gateway), "
        "and NoSQL metadata management (DynamoDB), with a responsive React frontend experience."
    )

    output_name = "Serverless_File_Upload_System_Documentation.docx"
    doc.save(output_name)
    print(f"Created: {output_name}")


if __name__ == "__main__":
    main()
